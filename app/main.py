

from imageBinarization import imageBinarizationAdaptive
from datetime import datetime
import zipfile
import os
from flask import Flask, request, render_template, url_for, send_from_directory, redirect
from werkzeug.utils import secure_filename

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt


app = Flask(__name__)
app.config['OUTPUT_FOLDER'] = 'outputs'  # 新增输出目录配置
app.config['UPLOAD_FOLDER'] = 'uploads'  # 新增上传目录配置


#methods参数用于指定允许的请求格式

#常规输入url的访问就是get方法
@app.route("/test",methods=['GET','POST'])
def test():
    return "Hello Test!"

# 可以在路径内以/<参数名>的形式指定参数，默认接收到的参数类型是string

'''#######################
以下为框架自带的转换器，可以置于参数前将接收的参数转化为对应类型
string 接受任何不包含斜杠的文本
int 接受正整数
float 接受正浮点数
path 接受包含斜杠的文本
########################'''


@app.route('/index')
def index():
    # redirect重定位（服务器向外部发起一个请求跳转）到一个url界面；
    # url_for给指定的函数构造 URL；
    # return redirect('/hello') 不建议这样做,将界面限死了
    return redirect(url_for('hello'))



# 欲实现url与视图函数的绑定，除了使用路由装饰器@app.route，我们还可以通过add_url_rule(rule,endpoint=None,view_func=None)方法，其中：
def test_url_rule():
    return '这是 test_url_rule 测试页面'
app.add_url_rule(rule='/test_url_rule',endpoint='test_url_rule',view_func=test_url_rule)






# 请求钩子before/after_request
# 想要在正常执行的代码的前、中、后时期，强行执行一段我们想要执行的功能代码，便要用到钩子函数——用特定装饰器装饰的函数。
# @app.before_request
# def before_request_a():
#     print('I am in before_request_a')
#

# @app.before_request
# def before_request_b():
#     print('I am in before_request_b')


# after_request：每一次请求之后都会调用；
# 该钩子函数表示每一次请求之后，可以执行某个特定功能的函数，这个函数接收response对象，所以执行完后必须归还response对象
# 执行的顺序是先绑定的后执行；
# 一般可以用于产生csrf_token验证码等场景；
# @app.after_request
# def after_request_a(response):
#     print('I am in after_request_a')
#     # 该装饰器接收response参数，运行完必须归还response，不然程序报错
#     return response
#
#
# @app.after_request
# def after_request_b(response):
#     print('I am in after_request_b')
#     return response
#
#



@app.route("/test_index", methods=['GET', 'POST'])
# url映射的函数，要传参则在上述route（路由）中添加参数申明
def test_index():
    if request.method == 'GET':
        # 想要html文件被该函数访问到，首先要创建一个templates文件，将html文件放入其中
        # 该文件夹需要被标记为模板文件夹，且模板语言设置为jinja2
        return render_template('index.html')
    # 此处欲发送post请求，需要在对应html文件的form表单中设置method为post
    elif request.method == 'POST':
        name = request.form.get('name')
        password = request.form.get('password')
        return name + " " + password















def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'jpg', 'png', 'jpeg'}




#  图片处理为Word文档
def create_word_file(processed_images):
    word_filename = os.path.join(create_date_subdirectory(app.config['OUTPUT_FOLDER']),'processed_images.docx')
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    for _, _, processed_img_url in processed_images:
        local_path = processed_img_url.replace('/', os.sep).split(app.config['OUTPUT_FOLDER'] + os.sep)[1]
        img_path = os.path.join(app.config['OUTPUT_FOLDER'], local_path)
        img = Image.open(img_path)
        img_width, img_height = img.size
        img_ratio = img_height / img_width
        doc_img_width = Inches(8.0)  # 80% of the page width minus margins
        doc_img_height = doc_img_width * img_ratio

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_path, width=doc_img_width, height=doc_img_height)
        paragraph.alignment = 1  # Center alignment
        paragraph.space_after = Pt(10)  # 10 pixels spacing after each image

    doc.save(word_filename)
    return word_filename






#     图片处理为PDF文件
def create_pdf_file(processed_images):
    pdf_filename = os.path.join(create_date_subdirectory(app.config['OUTPUT_FOLDER']), 'processed_images.pdf')
    c = canvas.Canvas(pdf_filename, pagesize=letter)
    width, height = letter
    margin_top_bottom = height * 0.07
    margin_left_right = width * 0.07
    pdf_img_width = width * 0.86
    y_offset = height - margin_top_bottom
    spacing = 10  # 图片之间的间隔距离

    for _, _, processed_img_url in processed_images:
        local_path = processed_img_url.replace('/', os.sep).split(app.config['OUTPUT_FOLDER'] + os.sep)[1]
        img_path = os.path.join(app.config['OUTPUT_FOLDER'], local_path)
        img = Image.open(img_path)
        img_width, img_height = img.size
        img_ratio = img_height / img_width
        pdf_img_height = pdf_img_width * img_ratio

        if y_offset - pdf_img_height < margin_top_bottom:
            c.showPage()
            y_offset = height - margin_top_bottom

        c.drawImage(img_path, margin_left_right, y_offset - pdf_img_height, width=pdf_img_width, height=pdf_img_height)
        y_offset -= (pdf_img_height + spacing)

    c.save()
    return pdf_filename



#  打包压缩过程
def create_date_subdirectory(base_folder):
    date_str = datetime.now().strftime('%Y-%m-%d')
    subdirectory = os.path.join(base_folder, date_str)
    if not os.path.exists(subdirectory):
        os.makedirs(subdirectory)
    return subdirectory

def create_zip_file(processed_images):
    zip_filename = secure_filename(processed_images[0][0].rsplit('.', 1)[0] + '.zip')
    zip_filepath = os.path.join(create_date_subdirectory(app.config['OUTPUT_FOLDER']), zip_filename)
    with zipfile.ZipFile(zip_filepath, 'w') as zipf:
        for _, _, processed_img in processed_images:
            # 获取文件的本地路径
            local_path = processed_img.replace('/', os.sep).split(app.config['OUTPUT_FOLDER'] + os.sep)[1]
            zipf.write(os.path.join(app.config['OUTPUT_FOLDER'], local_path), local_path)
    return zip_filepath.replace('\\', '/')

@app.route('/upload', methods=['GET', 'POST'])
def upload_and_process():
    if request.method == 'POST':
        if 'files' not in request.files:
            return redirect(request.url)
        files = request.files.getlist('files')
        if not files or all(file.filename == '' for file in files):
            return redirect(request.url)

        processed_images = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                upload_subdirectory = create_date_subdirectory(app.config['UPLOAD_FOLDER'])
                output_subdirectory = create_date_subdirectory(app.config['OUTPUT_FOLDER'])
                filepath = os.path.join(upload_subdirectory, filename)
                file.save(filepath)

                try:
                    output_filepath = os.path.join(output_subdirectory, filename)
                    imageBinarizationAdaptive(filepath, output_filepath)
                    # 手动拼接路径，避免 URL 编码问题
                    original_img_url = f"/uploads/{upload_subdirectory.replace(app.config['UPLOAD_FOLDER'] + os.sep, '').replace(os.sep, '/')}/{filename}"
                    processed_img_url = f"/outputs/{output_subdirectory.replace(app.config['OUTPUT_FOLDER'] + os.sep, '').replace(os.sep, '/')}/{filename}"
                    processed_images.append((filename, original_img_url, processed_img_url))
                except Exception as e:
                    return f"Error processing file: {str(e)}"

        # 生成 zip 文件
        if processed_images:
            zip_filename = create_zip_file(processed_images)
            pdf_filename = create_pdf_file(processed_images)
            word_filename = create_word_file(processed_images)
            return render_template('upload.html', processed_images=processed_images,
                                   zip_download_link='/' + zip_filename,
                                   pdf_download_link='/' + pdf_filename,
                                   word_download_link='/' + word_filename)

    return render_template('upload.html')

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/outputs/<path:filename>')
def uploaded_processed_file(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename)

@app.route('/download_zip/<filename>')
def download_zip(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)






if __name__ == '__main__':
    os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)  # 确保输出目录存在
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)  # 确保上传目录存在
    app.run(host='0.0.0.0', port=5000)
